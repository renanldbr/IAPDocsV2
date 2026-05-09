import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Utilitários ────────────────────────────────────────────────────────────────

MAPEAMENTO_DISCIPLINAS = {
    "FIS": "Física",
    "MAT": "Matemática",
    "BIO": "Biologia",
    "QUI": "Química",
    "CIE": "Ciências",
    "GEO": "Geografia",
    "HIS": "História",
    "LPT": "Língua Portuguesa",
    "ING": "Inglês",
    "ART": "Arte",
    "EFI": "Educação Física",
}


def limpar_texto(t):
    if not t:
        return ""
    bullets = r"[•●·▪\-*○◦⦿•‣⁃■□◊►▼]"
    t = re.sub(f"^{bullets}\\s*", "", t.strip())
    t = re.sub(f"\n{bullets}\\s*", "\n", t)
    t = re.sub(f"\\s+{bullets}\\s+", " ", t)
    return " ".join(t.split()).strip()


def set_cell_text(cell, text):
    """Limpa e define texto em uma célula preservando o parágrafo."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        cell.add_paragraph(text)


# ── Preenchedor EFAPE ──────────────────────────────────────────────────────────

def preencher_efape(doc, aulas, disciplina, etapa, serie, bimestre,
                    professor_nome, periodo, metodologias, recursos,
                    avaliacao, recuperacao):

    disciplina_completa = MAPEAMENTO_DISCIPLINAS.get(disciplina.upper(), disciplina)
    aulas_ordenadas = sorted(aulas, key=lambda x: int(x['numero']))

    re_ae_prefix = re.compile(r"^(AE|EA|AI)\d+\s*[-:]\s*", re.IGNORECASE)

    # ── Tabela 1: cabeçalho ──
    t1 = doc.tables[1]

    def append_to_cell(cell, label, value):
        if label.lower() in cell.text.lower():
            if cell.paragraphs:
                p = cell.paragraphs[0]
                run = p.add_run(f" {value}")
                run.bold = True

    append_to_cell(t1.cell(1, 0), "Professor(a):", professor_nome)
    append_to_cell(t1.cell(1, 1), "Ano/Série:", serie)
    append_to_cell(t1.cell(2, 0), "Componente Curricular:", disciplina_completa)
    t1.cell(2, 1).text = f"{bimestre}"

    for row in t1.rows:
        for cell in row.cells:
            if "período" in cell.text.lower():
                append_to_cell(cell, "Período de realização:", periodo)

    # ── Tabela 2: conteúdo pedagógico ──
    titulos = []
    conteudos_formatados = []
    objetivos_formatados = []
    mets_formatadas = []
    ae_unicos = []
    habilidades_unicas = []

    for i, aula in enumerate(aulas_ordenadas):
        prefix = f"Aula {aula['numero']}: "

        titulo_limpo = limpar_texto(aula['titulo'])
        titulos.append(f"{prefix}{titulo_limpo}")

        cont_limpo = limpar_texto(aula['conteudo']).replace("\n", "; ")
        conteudos_formatados.append(f"{prefix}{cont_limpo}")

        objs_limpos = "; ".join([limpar_texto(obj) for obj in aula['objetivos_aprendizagem']])
        objetivos_formatados.append(f"{prefix}{objs_limpos}")

        met_texto = metodologias[i] if i < len(metodologias) else ""
        mets_formatadas.append(f"{prefix}{met_texto}")

        ae_raw = aula['aprendizagem_essencial'].strip()
        ae_clean = re_ae_prefix.sub("", ae_raw).strip()
        ae_norm = " ".join(ae_clean.strip(" .").split()).lower()
        if not any(" ".join(x.strip(" .").split()).lower() == ae_norm for x in ae_unicos):
            ae_unicos.append(ae_clean)

        hab = aula['habilidade'].strip()
        if not any(
            " ".join(x.strip(" .").split()).lower() == " ".join(hab.strip(" .").split()).lower()
            for x in habilidades_unicas
        ):
            habilidades_unicas.append(hab)

    ae_formatado = [f"AE{idx+1}: {ae_val}" for idx, ae_val in enumerate(ae_unicos)]

    t2 = doc.tables[2]
    t2.cell(1, 0).text = "\n".join(titulos)
    t2.cell(3, 0).text = "\n".join(ae_formatado)
    t2.cell(5, 0).text = "\n".join(habilidades_unicas)
    t2.cell(7, 0).text = "\n".join(conteudos_formatados)
    t2.cell(9, 0).text = "\n".join(objetivos_formatados)

    if len(t2.rows) > 11:
        t2.cell(11, 0).text = "\n".join(mets_formatadas)
    if len(t2.rows) > 13:
        t2.cell(13, 0).text = recursos
    if len(t2.rows) > 15:
        t2.cell(15, 0).text = avaliacao
    if len(t2.rows) > 19:
        t2.cell(19, 0).text = recuperacao


# ── Preenchedor Guia de Aprendizagem ─────────────────────────────────────────

def preencher_guia_aprendizagem(doc, aulas, disciplina, etapa, serie, bimestre,
                               professor_nome, metodologias, ambientes,
                               avaliacao, fontes, justificativa, aproximacao):

    disciplina_completa = MAPEAMENTO_DISCIPLINAS.get(disciplina.upper(), disciplina)
    aulas_ordenadas = sorted(aulas, key=lambda x: int(x['numero']))
    t = doc.tables[0]

    # ── Linha 1: valores do cabeçalho (linha 0 tem os labels) ──
    t.cell(1, 0).text = professor_nome
    t.cell(1, 2).text = disciplina_completa
    t.cell(1, 4).text = serie
    t.cell(1, 5).text = bimestre

    # ── Linha 3: Justificativa (editável) ──
    t.cell(3, 0).text = justificativa

    # ── Linha 5: Aproximação com a realidade (editável) ──
    t.cell(5, 0).text = aproximacao

    # ── Linha 7: Título | Conteúdos | Objetivos ──
    titulos = []
    conteudos = []
    objetivos = []
    for aula in aulas_ordenadas:
        prefix = f"Aula {aula['numero']}: "
        titulos.append(f"{prefix}{limpar_texto(aula['titulo'])}")
        cont = limpar_texto(aula['conteudo']).replace('\n', '; ')
        conteudos.append(f"{prefix}{cont}")
        objs = '; '.join([limpar_texto(o) for o in aula['objetivos_aprendizagem']])
        objetivos.append(f"{prefix}{objs}")

    t.cell(7, 0).text = '\n'.join(titulos)
    t.cell(7, 1).text = '\n'.join(conteudos)
    t.cell(7, 3).text = '\n'.join(objetivos)

    # ── Linha 9: Metodologias | Ambientes de Aprendizagem ──
    mets_formatadas = []
    for i, aula in enumerate(aulas_ordenadas):
        met_texto = metodologias[i] if i < len(metodologias) else ''
        mets_formatadas.append(f"Aula {aula['numero']}: {met_texto}")

    t.cell(9, 0).text = '\n'.join(mets_formatadas)
    t.cell(9, 3).text = ambientes

    # ── Linha 11: Critérios de Avaliação ──
    t.cell(11, 0).text = avaliacao

    # ── Linha 13: Fontes de pesquisa ──
    t.cell(13, 0).text = fontes


# ── Preenchedor PAV2 ───────────────────────────────────────────────────────────

def preencher_pav2(doc, aulas, disciplina, etapa, serie, bimestre,
                   professor_nome, periodo, metodologias, recursos,
                   avaliacao, recuperacao):

    disciplina_completa = MAPEAMENTO_DISCIPLINAS.get(disciplina.upper(), disciplina)
    aulas_ordenadas = sorted(aulas, key=lambda x: int(x['numero']))

    t = doc.tables[0]

    # ── Linha 0: cabeçalho com info do professor ──
    header_cell = t.cell(0, 0)
    header_cell.text = (
        f"Nome do professor: {professor_nome}\n"
        f"Componente Curricular: {disciplina_completa}\n"
        f"Ano/ Série: {serie} | {bimestre} | Período: {periodo}"
    )

    # ── Linhas de aulas (índice 2 em diante) ──
    # Colunas PAV2:
    # 0: Aula / Data
    # 1: Objetivos de aprendizagem
    # 2: Habilidade
    # 3: Competências socioemocionais
    # 4: Metodologia
    # 5: Desenvolvimento da aula
    # 6-7: Recursos pedagógicos
    # 8: Avaliação

    DATA_START_ROW = 2  # linha 1 é o cabeçalho das colunas

    for i, aula in enumerate(aulas_ordenadas):
        row_idx = DATA_START_ROW + i
        # Se não houver linhas suficientes no template, para
        if row_idx >= len(t.rows):
            break

        row = t.rows[row_idx]

        objs = "; ".join([limpar_texto(obj) for obj in aula['objetivos_aprendizagem']])
        met_texto = metodologias[i] if i < len(metodologias) else ""
        conteudo = limpar_texto(aula['conteudo'])

        # Aula / Data
        set_cell_text(row.cells[0], f"Aula {aula['numero']}")
        # Objetivos
        set_cell_text(row.cells[1], objs)
        # Habilidade
        set_cell_text(row.cells[2], limpar_texto(aula['habilidade']))
        # Competências socioemocionais (deixa em branco para o professor preencher)
        set_cell_text(row.cells[3], "")
        # Metodologia
        set_cell_text(row.cells[4], met_texto)
        # Desenvolvimento / conteúdo
        set_cell_text(row.cells[5], conteudo)
        # Recursos pedagógicos
        set_cell_text(row.cells[6], recursos)
        # Recursos pedagógicos (col duplicada)
        set_cell_text(row.cells[7], "")
        # Avaliação
        set_cell_text(row.cells[8], avaliacao)


# ── Função principal ───────────────────────────────────────────────────────────

def preencher_documento(modelo_path, modelo_tipo, aulas, disciplina, etapa,
                        serie, bimestre, professor_nome, periodo, metodologias,
                        recursos, avaliacao, recuperacao,
                        # Campos extras para o Guia de Aprendizagem
                        ambientes="", fontes="", justificativa="", aproximacao=""):
    """
    Preenche o documento Word com os dados das aulas selecionadas.
    modelo_tipo: 'EFAPE', 'PAV2' ou 'GUIA'
    Retorna o caminho do arquivo gerado.
    """
    doc = Document(modelo_path)

    if modelo_tipo == "PAV2":
        preencher_pav2(
            doc, aulas, disciplina, etapa, serie, bimestre,
            professor_nome, periodo, metodologias, recursos, avaliacao, recuperacao
        )
    elif modelo_tipo == "GUIA":
        preencher_guia_aprendizagem(
            doc, aulas, disciplina, etapa, serie, bimestre,
            professor_nome, metodologias, ambientes,
            avaliacao, fontes, justificativa, aproximacao
        )
    else:
        preencher_efape(
            doc, aulas, disciplina, etapa, serie, bimestre,
            professor_nome, periodo, metodologias, recursos, avaliacao, recuperacao
        )

    disciplina_completa = MAPEAMENTO_DISCIPLINAS.get(disciplina.upper(), disciplina)
    safe_disciplina = disciplina_completa.replace("/", "-").replace(" ", "_")
    safe_serie = serie.replace("/", "-").replace(" ", "_")
    safe_bimestre = bimestre.replace("/", "-").replace(" ", "_")

    os.makedirs("./docOrientadores/Saida", exist_ok=True)
    saida_path = (
        f"./docOrientadores/Saida/Plano_{modelo_tipo}_{safe_disciplina}_{safe_serie}_{safe_bimestre}.docx"
    )
    doc.save(saida_path)
    return saida_path
